from __future__ import annotations

import logging
import os
from io import BytesIO

log = logging.getLogger("pptx-service.parsing")

MAX_CHARS = 40000


class UnsupportedFormat(ValueError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext((filename or "").lower())[1]
    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext in (".txt", ".md", ".markdown", ".rst"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise UnsupportedFormat(f"unsupported file extension: {ext!r}")

    text = text.strip()
    if len(text) > MAX_CHARS:
        log.warning("truncating %s from %d to %d chars", filename, len(text), MAX_CHARS)
        text = text[:MAX_CHARS]
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:
            log.warning("pdf page extract failed: %s", e)
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
